# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

FONT_NAME = "맑은 고딕"

def set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, size=16, color=(0x1F, 0x4E, 0x79), space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_note(doc, text, size=10, italic=True, color=(0x66, 0x66, 0x66)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, color=color)
    run.italic = italic
    return p

def add_body(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        row.height = Cm(0.9)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_info_table(doc, rows, label_width=4.5, value_width=11.0):
    table = doc.add_table(rows=len(rows), cols=2)
    style_table(table)
    table.columns[0].width = Cm(label_width)
    table.columns[1].width = Cm(value_width)
    for i, (label, value) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c0.text = ""
        run = c0.paragraphs[0].add_run(label)
        set_font(run, size=10.5, bold=True)
        c0.paragraphs[0].paragraph_format.space_after = Pt(0)

        c1 = table.rows[i].cells[1]
        c1.text = ""
        run2 = c1.paragraphs[0].add_run(value or "")
        set_font(run2, size=10.5)
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)
    return table

def add_essay_item(doc, number, question):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {question}")
    set_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    cell = box.rows[0].cells[0]
    cell.height = Cm(2.4)
    for _ in range(3):
        cell.add_paragraph("")
    tr = box.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.makeelement(qn('w:trHeight'), {qn('w:val'): "1400", qn('w:hRule'): "atLeast"})
    trPr.append(trHeight)

def add_checkbox(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("☐  ")
    set_font(run, size=11, bold=True)
    run2 = p.add_run(text)
    set_font(run2, size=11)


doc = Document()

# 기본 여백/폰트
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

normal = doc.styles["Normal"]
normal.font.name = FONT_NAME
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# 제목
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("발스 클랜 지원서")
set_font(run, size=24, bold=True, color=(0x1F, 0x4E, 0x79))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("실력과 매너를 함께 보는 클랜입니다")
set_font(run, size=11, color=(0x66, 0x66, 0x66))
subtitle.italic = True

doc.add_paragraph()

# 1. 기본 정보
add_heading(doc, "1. 기본 정보")
add_info_table(doc, [
    ("실명", ""),
    ("닉네임(게임 내)", ""),
    ("생년월일", ""),
    ("연락처(카카오톡 ID)", ""),
    ("현재/과거 소속 클랜", "(없으면 \"없음\")"),
])

# 2. 실력 정보
add_heading(doc, "2. 실력 정보  ★가장 중요★")
add_info_table(doc, [
    ("최고 레이팅\n(역대 달성 최고치)", ""),
    ("최고 레이팅 달성 시점", "예) 2026년 3월"),
    ("인증 스크린샷", "최고 레이팅 화면 캡처 첨부 필수 (미첨부 시 심사 보류)"),
    ("현재 레이팅", ""),
    ("주 포지션 / 플레이스타일", "예) 공격형, 점유율, 역습형 등"),
    ("주력 사용팀", ""),
])
add_note(doc, "※ 실력 평가는 자기 신고 최고 레이팅 + 스크린샷 대조로 진행합니다. 레이팅 허위 기재가 확인될 경우 지원이 즉시 반려되며, 가입 이후 발견 시 클랜에서 제명될 수 있습니다.")

# 3. 활동 가능 정보
add_heading(doc, "3. 활동 가능 정보")
add_info_table(doc, [
    ("리그 참여 가능 여부", "예 / 아니오"),
    ("내전 참여 가능 요일", "예) 월·수·금·주말"),
    ("주로 접속 가능한 시간대", "예) 평일 21~24시"),
    ("예상 주간 참여 빈도", "예) 주 3~4회"),
])

# 4. 매너 및 인성
add_heading(doc, "4. 매너 및 인성 (실력만큼 중요합니다)")
add_essay_item(doc, 1, "이전 클랜에서 활동한 적이 있다면, 탈퇴/추방 사유를 솔직하게 적어주세요. (없으면 \"해당 없음\")")
add_essay_item(doc, 2, "경기 중 욕설, 고의 트롤링, 비매너 플레이(고의 지연·도배 등)에 대해 어떻게 생각하시나요?")
add_essay_item(doc, 3, "내전/리그에서 판정(스코어 입력, 팀 배정 등)에 이견이 생겼을 때 본인은 어떻게 대응하는 편인가요?")
add_essay_item(doc, 4, "발스 클랜에 지원하는 이유를 간단히 적어주세요.")

# 5. 서약
add_heading(doc, "5. 서약")
add_body(doc, "아래 항목에 모두 동의해야 지원이 접수됩니다.")
doc.add_paragraph()
add_checkbox(doc, "위에 기재한 최고 레이팅 정보는 사실이며, 허위 기재 시 가입 취소·제명에 동의합니다.")
add_checkbox(doc, "경기 중 욕설, 고의 트롤링, 비매너 행위를 하지 않을 것을 서약합니다.")
add_checkbox(doc, "리그·내전 참여 시 스코어 입력 등 클랜 앱 이용 규칙을 성실히 따르겠습니다.")
add_checkbox(doc, "무단 장기 미접속(잠수) 시 클랜 정리 대상이 될 수 있음에 동의합니다.")

doc.add_paragraph()
sign_table = add_info_table(doc, [
    ("지원일자", ""),
    ("서명(닉네임)", ""),
])

doc.save(r"c:\cursor\balsleague\발스클랜_지원서.docx")
print("saved")
